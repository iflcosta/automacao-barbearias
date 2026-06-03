import os
import subprocess
import sys

# Garante que as bibliotecas necessárias estão instaladas
def install_libraries():
    try:
        import docx
        import reportlab
    except ImportError:
        print("Instalando dependências (python-docx e reportlab)...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "reportlab"])

install_libraries()

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

proposal_title = "Proposta Comercial: Ecossistema de Automação para Barbearias"

proposal_content = [
    ("O problema do mercado", 
     "Barbeiros e donos de barbearias são profissionais extremamente operacionais. Eles perdem clientes diariamente por dois motivos principais:\n\n"
     "1. Demora no Atendimento no WhatsApp: O cliente quer agendar um horário rápido e, se a barbearia demora 15 minutos para responder, ele agenda com o concorrente.\n"
     "2. Falta de Consistência no Marketing: Não têm tempo para postar no Instagram, criar páginas web modernas, gerenciar anúncios locais ou fazer pós-venda para fidelizar quem já é cliente."),
    
    ("A Solução: Ecossistema de Automação com IA", 
     "Uma plataforma 'tudo em um' (All-in-One) que assume toda a operação digital da barbearia de forma automatizada, gerando mais agendamentos e retendo clientes com custo operacional quase zero. Toda a estrutura roda em servidores de baixo custo na nuvem."),
    
    ("Pilar 1: Assistente de Agendamento Inteligente (WhatsApp 24/7)", 
     "Status: Protótipo Desenvolvido e Validado Localmente.\n\n"
     "O que faz: Um assistente virtual no WhatsApp conectado à agenda do estabelecimento. Ele conversa de forma humana, entende gírias e horários relativos (ex: 'amanhã de tarde'), consulta os horários livres no Google Calendar em tempo real, efetua a reserva automaticamente e lida com cancelamentos."),
    
    ("Pilar 2: CRM Ativo de Retenção e Pós-Venda (WhatsApp Proativo)", 
     "O que faz: Varre o banco de dados e identifica clientes que não cortam o cabelo há mais de 25 dias. O sistema gera uma mensagem personalizada e amigável lembrando o cliente de renovar o visual antes do final de semana.\n\n"
     "Diferencial: Aumenta a frequência de visitas e o faturamento geral de forma previsível."),
    
    ("Pilar 3: Geração Automatizada de Landing Pages (Sites)", 
     "O que faz: A partir de um questionário rápido, nossos agentes de IA criam o copywriting do site, montam o design moderno responsivo e fazem o deploy do site na nuvem automaticamente, otimizando o SEO Local."),
    
    ("Pilar 4: Criador Automático de Conteúdo para Mídias Sociais", 
     "O que faz: Agentes de IA planejam um cronograma mensal de posts, escrevem legendas engajadoras e geram imagens personalizadas (usando modelos de geração de imagem locais de alta performance) com a marca da barbearia."),
    
    ("Pilar 5: Gestor de Tráfego Pago & Otimizador de Anúncios", 
     "O que faz: Gera copies de alta conversão para anúncios locais no Google/Instagram. No final do mês, a IA lê os dados das campanhas de anúncios e gera um relatório simplificado e intuitivo sobre o retorno financeiro."),
    
    ("Viabilidade Técnica e Escalabilidade", 
     "O grande diferencial técnico é a eficiência de custos:\n"
     "- Arquitetura Leve (Multi-Tenant): Desenvolvido em TypeScript e SQLite, consumindo menos de 80MB de RAM por backend, ideal para rodar múltiplos clientes em uma VPS de 1GB.\n"
     "- IA de Baixo Custo: Uso da API da Groq para processamento de texto e workstation local para renderizar as imagens semanais, maximizando a margem de lucro do ecossistema.")
]

def generate_docx():
    print("Gerando arquivo DOCX...")
    doc = Document()
    
    # Configura fontes e estilo geral do documento
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Adiciona título principal com cor customizada (azul escuro premium)
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(proposal_title)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(26, 82, 118)
    
    doc.add_paragraph("Apresentação de Projeto e Parcerias Comerciais\n")
    
    for title, text in proposal_content:
        h = doc.add_paragraph()
        h_run = h.add_run(title)
        h_run.bold = True
        h_run.font.size = Pt(13)
        h_run.font.color.rgb = RGBColor(41, 128, 185)
        
        doc.add_paragraph(text)
        doc.add_paragraph()
        
    doc.save("Proposta_Comercial_Automação_Barbearias.docx")
    print("Arquivo DOCX gerado com sucesso!")

def generate_pdf():
    print("Gerando arquivo PDF otimizado para celular...")
    # Margens menores facilitam a leitura em telas pequenas
    doc = SimpleDocTemplate(
        "Proposta_Comercial_Automação_Barbearias.pdf",
        pagesize=letter,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36
    )
    
    styles = getSampleStyleSheet()
    
    # Estilos customizados
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#1A5276'),
        spaceAfter=15
    )
    
    heading_style = ParagraphStyle(
        'HeadingStyle',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=16,
        textColor=colors.HexColor('#2980B9'),
        spaceBefore=12,
        spaceAfter=6
    )
    
    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#2C3E50'),
        spaceAfter=10
    )

    story = []
    story.append(Paragraph(proposal_title, title_style))
    story.append(Paragraph("<b>Apresentação do Ecossistema de Inteligência Artificial Local e Nuvem</b>", body_style))
    story.append(Spacer(1, 10))
    
    for title, text in proposal_content:
        story.append(Paragraph(title, heading_style))
        # Preserva quebras de linha no PDF
        formatted_text = text.replace('\n', '<br/>')
        story.append(Paragraph(formatted_text, body_style))
        
    doc.build(story)
    print("Arquivo PDF gerado com sucesso!")

if __name__ == "__main__":
    generate_docx()
    generate_pdf()
